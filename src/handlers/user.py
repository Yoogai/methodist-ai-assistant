import logging
import io
import textwrap
import base64
import re
from typing import Optional, List, Tuple, Dict, Any

from aiogram import Router, F, Bot
from aiogram.filters import CommandStart, StateFilter
from aiogram.fsm.context import FSMContext
from aiogram.fsm.state import State, StatesGroup
from aiogram.types import Message, CallbackQuery, FSInputFile, BufferedInputFile
from aiogram.exceptions import TelegramBadRequest
from html import escape
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Сервисы
from src.services.rag_engine import RagEngine
from src.services.yandex_gpt import YandexGPTService
from src.services.web_search_service import YandexWebSearchService
from src.services.speech_service import YandexSpeechKitService
from src.services.file_search_service import FileSearchService
from src.services.database import db
from src.services.ocr_service import YandexOCRService

# Клавиатуры
from src.keyboards.builders import (
    create_smart_keyboard,
    get_main_menu_keyboard,
    create_settings_keyboard,
    create_creative_keyboard,
    create_recognition_keyboard
)
from src.config import ADMIN_ID, PDF_DIR

logger = logging.getLogger(__name__)
router = Router()

# Инициализация
rag_service = RagEngine()
gpt_service = YandexGPTService()
web_search_service = YandexWebSearchService()
speech_service = YandexSpeechKitService()
file_search_service = FileSearchService()
ocr_service = YandexOCRService()

# --- Константы ---
STARTUP_SUGGESTIONS = ["Об НМО НБ РА", "Правила оформления методички", "О комплектовании фондов"]
FILE_REQUEST_TRIGGERS = ["скинь", "дай", "пришли", "отправь", "файл", "документ", "график", "список"]
MAX_AUDIO_SIZE = 1024 * 1024

# --- Промпты ---

SYSTEM_PROMPT = """Ты — «Методист НБ РА», ведущий эксперт-консультант. 
Твоя задача — дать исчерпывающий и структурированный ответ, используя ТОЛЬКО предоставленный контекст.

ПРАВИЛА:
1. Если в контексте нет ответа, цифр или конкретных фактов, прямо пиши: «В предоставленных методических материалах данная информация отсутствует». 
2. Категорически запрещено выдумывать номера приказов, даты или фамилии.
3. Используй Markdown-подобное форматирование, но для Telegram используй теги <b>жирный</b>, <i>курсив</i>.
4. Тон общения: официально-деловой, но не сухой. Избегай канцеляризмов."""

# Обновленный промпт для общения
CHIT_CHAT_PROMPT = """Ты — «Цифровой помощник НМО НБ РА», дружелюбный и компетентный ассистент. 
Твоя цель — помогать сотрудникам библиотек, избегая сложных канцелярских фраз. Будь живым собеседником.

СЛЕДУЙ ЭТИМ СЦЕНАРИЯМ:

1. **Если это приветствие** (Привет, Здравствуйте) -> Поздоровайся тепло, представься и предложи помощь.

2. **Если спрашивают "Что ты умеешь?", "Твои функции", "О чем рассказать?"** ->
   Ответь: "Я готов проконсультировать вас по следующим темам:"
   Затем выведи список (используй теги <b></b> для заголовков):
   • 📚 <b>Комплектование и учёт</b> библиотечных фондов.
   • 📝 <b>Оформление методических пособий</b> и изданий.
   • 📊 <b>Статистический учёт</b> (форма 6-НК и др.).
   • 🏛 <b>Работа научно-методического отдела</b> Национальной библиотеки РА.
   • 📰 <b>Библиографические обзоры</b> и списки литературы.

   Закончи фразой: "Просто задайте вопрос в свободной форме."

3. **Если это благодарность** (Спасибо, Благодарю) -> Ответь: "Всегда пожалуйста! Рад быть полезным."

4. **Если вопрос не по теме** (погода, новости, рецепты) -> Вежливо скажи: "К сожалению, пока я не владею информацией по этому вопросу. Но я быстро учусь! Попробуйте спросить что-нибудь о библиотечном деле."
"""

OCR_CLEANUP_PROMPT = """Ты — корректор. Исправь текст OCR: 
1. Соедини слова, разорванные переносом. 
2. Удали номера страниц и колонтитулы. 
3. Сделай текст цельным повествованием. 
Верни ТОЛЬКО чистый текст."""

VLM_COMPLEX_PROMPT = """Проанализируй изображение. Извлеки ВЕСЬ текст и таблицы.
ВНИМАНИЕ: Верни ТОЛЬКО содержание документа. 
КАТЕГОРИЧЕСКИ ЗАПРЕЩЕНО писать вводные фразы вроде "Вот текст", "Результат анализа".
Сразу начинай с заголовка или текста документа.
Таблицы оформляй через Markdown (символ |)."""

VLM_DESCRIBE_PROMPT = "Опиши детально, что изображено на этой фотографии. Используй <b> для акцентов."
IDEA_PROMPT = "Ты — аналитик. Структурируй сообщение пользователя (идею или баг-репорт) с помощью HTML-тегов <b>."

POST_PROMPT = "Напиши яркий пост для соцсетей. Используй <b> для заголовков."
PRESS_RELEASE_PROMPT = "Напиши официальный пресс-релиз. Используй <b> для важных данных."
ANNOUNCEMENT_PROMPT = "Напиши анонс мероприятия. Используй <b> для даты и места."
CUSTOM_CREATIVE_PROMPT = "Помоги создать текст, задавая вопросы. Используй <b> для выделения сути."


class DialogStates(StatesGroup):
    main = State()
    web_search = State()
    feedback = State()
    settings = State()
    idea_mode = State()
    creative_mode = State()
    creative_post = State()
    creative_release = State()
    creative_announcement = State()
    creative_custom = State()
    recognition_mode = State()


# --- Утилиты ---

def create_formatted_docx(md_text: str, title: str = "Распознанный документ") -> io.BytesIO:
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    clean_text = re.sub(r'^(Вот|Here is|Результат|Analysis|Извлеченный|Ниже).*?[:\n]', '', md_text,
                        flags=re.IGNORECASE | re.DOTALL).strip()
    doc.add_heading(title, 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    lines = clean_text.split('\n')
    table_data = []
    in_table = False
    for line in lines:
        line = line.strip()
        if not line: continue
        if '|' in line:
            if '---' in line: continue
            cells = [c.strip() for c in line.split('|') if c.strip()]
            if cells:
                table_data.append(cells)
                in_table = True
            continue
        else:
            if in_table and table_data:
                try:
                    table = doc.add_table(rows=len(table_data), cols=max(len(r) for r in table_data))
                    table.style = 'Table Grid'
                    for r_idx, row_content in enumerate(table_data):
                        for c_idx, cell_text in enumerate(row_content):
                            table.cell(r_idx, c_idx).text = cell_text
                except:
                    pass
                doc.add_paragraph("")
                table_data = []
                in_table = False
            if line.startswith('#'):
                doc.add_heading(line.replace('#', '').strip(), level=min(line.count('#'), 9))
            else:
                p = doc.add_paragraph()
                parts = re.split(r'(\*\*.*?\*\*)', line)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part.replace('**', ''))
                        run.bold = True
                    else:
                        p.add_run(part)
    if in_table and table_data:
        try:
            table = doc.add_table(rows=len(table_data), cols=max(len(r) for r in table_data))
            table.style = 'Table Grid'
            for r_idx, row_content in enumerate(table_data):
                for c_idx, cell_text in enumerate(row_content):
                    table.cell(r_idx, c_idx).text = cell_text
        except:
            pass
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def clean_html_for_telegram(text: str) -> str:
    """Очищает HTML от запрещенных тегов и исправляет разметку."""
    # 1. Удаляем служебные заголовки
    text = re.sub(r'<!DOCTYPE.*?>', '', text, flags=re.IGNORECASE | re.DOTALL)

    # 2. Заменяем <br> на переносы строк (ВАЖНОЕ ИСПРАВЛЕНИЕ)
    text = re.sub(r'<br\s*/?>', '\n', text, flags=re.IGNORECASE)

    # 3. Заменяем параграфы и дивы на переносы
    text = re.sub(r'</?(p|div).*?>', '\n', text, flags=re.IGNORECASE)

    # 4. Заменяем таблицы на псевдо-таблицы
    text = re.sub(r'</?(table|tbody|thead|tr).*?>', '\n', text, flags=re.IGNORECASE)
    text = re.sub(r'<td.*?>', ' ', text, flags=re.IGNORECASE)
    text = re.sub(r'</td>', ' | ', text, flags=re.IGNORECASE)
    text = re.sub(r'<th.*?>', ' ', text, flags=re.IGNORECASE)
    text = re.sub(r'</th>', ' | ', text, flags=re.IGNORECASE)

    # 5. Очистка от Markdown-звездочек (если ИИ их оставил)
    text = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', text)

    # 6. Удаляем все остальные теги, кроме разрешенных Telegram
    # Разрешенные: b, strong, i, em, u, ins, s, strike, del, a, code, pre
    # Но проще просто экранировать всё, что не похоже на разрешенные,
    # однако мы уже полагаемся на то, что ИИ пишет <b>.
    # Если проскочит <script> или <style>, Telegram выдаст ошибку.
    # Для надежности можно вырезать всё, что не в белом списке, но пока ограничимся заменой BR.

    return text.strip()

async def send_split_message(message: Message, text: str, reply_markup=None):
    text = clean_html_for_telegram(text)
    if len(text) <= 4096:
        try:
            await message.answer(text, reply_markup=reply_markup, parse_mode="HTML")
        except:
            await message.answer(escape(text), reply_markup=reply_markup)
    else:
        chunks = textwrap.wrap(text, width=4000, replace_whitespace=False, drop_whitespace=False)
        for i, chunk in enumerate(chunks):
            is_last = (i == len(chunks) - 1)
            try:
                await message.answer(chunk, reply_markup=reply_markup if is_last else None, parse_mode="HTML")
            except:
                await message.answer(escape(chunk), reply_markup=reply_markup if is_last else None)


async def encode_image_to_base64(bot: Bot, photo_id: str) -> str:
    file_info = await bot.get_file(photo_id)
    photo_bytes = io.BytesIO()
    await bot.download_file(file_info.file_path, photo_bytes)
    return base64.b64encode(photo_bytes.getvalue()).decode('utf-8')


async def try_send_file(message: Message, bot: Bot) -> bool:
    if not message.text: return False
    user_text = message.text.lower()
    if any(trigger in user_text for trigger in FILE_REQUEST_TRIGGERS):
        file_data = file_search_service.find_file(user_text)
        if file_data:
            file_path = file_search_service.get_full_path(file_data["filename"])
            if file_path.exists():
                await bot.send_chat_action(message.chat.id, "upload_document")
                await message.reply_document(FSInputFile(file_path), caption=f"Вот файл: <b>{file_data['title']}</b>",
                                             parse_mode="HTML")
                return True
    return False


# --- ВАЖНАЯ ФУНКЦИЯ ДЛЯ ОПРЕДЕЛЕНИЯ "БОЛТОВНИ" ---
def is_small_talk(text: str) -> bool:
    """Проверяет, является ли сообщение приветствием или вопросом о возможностях."""
    text_lower = text.lower().strip()

    # Прямые совпадения и фразы
    triggers = [
        "привет", "здравствуй", "добрый день", "добрый вечер", "хай",
        "кто ты", "что ты", "что умеешь", "твои функции", "о чем рассказать",
        "помощь", "справка", "спасибо", "благодарю", "пока"
    ]

    # Если сообщение короткое (меньше 5 слов) и содержит триггер
    if len(text_lower.split()) < 6:
        for trigger in triggers:
            if trigger in text_lower:
                return True
    return False


# --- ЯДРО ЛОГИКИ ---

async def get_ai_response(state: FSMContext, user_id: int, user_text: str) -> Tuple[
    str, List[str], Optional[str], Optional[Dict[str, Any]]]:
    user_data = db.get_user(user_id)
    full_name = user_data.get("full_name") or user_data.get("first_name") or "Коллега"
    fsm_data = await state.get_data()
    history = fsm_data.get("history", [])
    recognized_context = fsm_data.get("last_recognized_text", "")

    context = ""
    metadata = None
    pdf_slug = None
    prompt = CHIT_CHAT_PROMPT  # По умолчанию считаем это болтовней

    # Логика выбора режима (Small Talk vs RAG)
    if is_small_talk(user_text) and not recognized_context:
        # Если это приветствие и нет контекста документа -> Отвечаем через CHIT_CHAT_PROMPT без поиска
        logger.info(f"Detected Small Talk: {user_text}")
        context = ""
    else:
        # Иначе ищем в базе
        context, metadata = rag_service.search(user_text)
        pdf_slug = metadata.get("slug") if metadata else None

        if context:
            prompt = SYSTEM_PROMPT
        else:
            # Если в базе ничего нет, но это не явный Small Talk - всё равно используем CHIT_CHAT
            # но промпт там настроен на вежливый отказ в п.4
            pass

    # Если есть контекст из фото, он приоритетнее
    full_context = f"КОНТЕКСТ ИЗ ФОТО:\n{recognized_context}\n\nБАЗА ЗНАНИЙ:\n{context}" if recognized_context else context

    # Если контекста вообще нет и это не Small Talk, бот ответит вежливым отказом (п.4 CHIT_CHAT_PROMPT)

    res = gpt_service.generate_response(prompt, user_text, full_context, history, full_name)
    ai_text = res.get("text", "Ошибка.")
    suggestions = res.get("suggestions", [])

    new_history = history + [{"role": "user", "text": user_text}, {"role": "assistant", "text": ai_text}]
    await state.update_data(history=new_history[-6:], last_query=user_text, last_suggestions=suggestions)
    return ai_text, suggestions, pdf_slug, metadata


# --- Хендлеры ---

@router.message(CommandStart())
async def command_start_handler(message: Message, state: FSMContext):
    db.add_user(message.from_user.id, message.from_user.username, message.from_user.full_name)
    await state.set_state(DialogStates.main)
    await state.update_data(history=[], settings={"voice_mode": "text_to_text"}, last_recognized_text="")
    await message.answer(f"Здравствуйте, <b>{escape(message.from_user.full_name)}</b>!",
                         reply_markup=get_main_menu_keyboard(), parse_mode="HTML")
    await message.answer("Я — ваш цифровой методист. Чем могу помочь?",
                         reply_markup=create_smart_keyboard(STARTUP_SUGGESTIONS, None), parse_mode="HTML")


@router.message(F.text == "⚙️ Параметры")
async def settings_handler(message: Message, state: FSMContext):
    await state.set_state(DialogStates.settings)
    fsm_data = await state.get_data()
    current_settings = fsm_data.get("settings", {"voice_mode": "text_to_text"})
    await message.answer("Настройте параметры работы бота:", reply_markup=create_settings_keyboard(current_settings))


@router.message(F.text == "💡 Есть идея")
async def idea_start(message: Message, state: FSMContext):
    await state.set_state(DialogStates.idea_mode)
    await message.answer("Опишите вашу идею или замечание. Я передам это разработчику.")


@router.message(F.text == "✍️ Написать нам")
async def feedback_start(message: Message, state: FSMContext):
    await state.set_state(DialogStates.feedback)
    await message.answer("Напишите ваше сообщение методисту:")


@router.message(F.text == "🌐 Поиск в сети")
async def web_search_handler(message: Message, state: FSMContext):
    await state.set_state(DialogStates.web_search)
    await message.answer("Введите ваш запрос для поиска в интернете:")


@router.message(StateFilter(DialogStates.idea_mode))
async def process_idea(message: Message, state: FSMContext, bot: Bot):
    res = gpt_service.generate_response(IDEA_PROMPT, message.text)
    formatted = res.get("text", message.text)
    report = f"💡 <b>ИДЕЯ/БАГ</b>\n👤 От: {escape(message.from_user.full_name)}\n🆔 ID: <code>{message.from_user.id}</code>\n---\n{clean_html_for_telegram(formatted)}"
    await bot.send_message(ADMIN_ID, report, parse_mode="HTML")
    await message.answer("✅ Спасибо! Сообщение передано.", reply_markup=get_main_menu_keyboard())
    await state.set_state(DialogStates.main)


@router.message(StateFilter(DialogStates.feedback))
async def process_feedback(message: Message, bot: Bot, state: FSMContext):
    admin_msg = f"✉️ <b>СООБЩЕНИЕ</b>\n👤 От: {escape(message.from_user.full_name)}\n🆔 ID: <code>{message.from_user.id}</code>\n---\n{escape(message.text)}"
    await bot.send_message(ADMIN_ID, admin_msg, parse_mode="HTML")
    await message.answer("✅ Доставлено.", reply_markup=get_main_menu_keyboard())
    await state.set_state(DialogStates.main)


@router.message(StateFilter(DialogStates.web_search))
async def process_web_search(message: Message, state: FSMContext):
    status_msg = await message.answer("🌐 Ищу...")
    try:
        res = web_search_service.generate_web_response(message.text)
        if res and isinstance(res, list) and res[0].get("message"):
            ans = res[0]["message"]["content"]
            await status_msg.delete()
            await send_split_message(message, f"<b>Результат поиска:</b>\n\n{ans}")
        else:
            await status_msg.edit_text("😕 Ничего не найдено.")
    except Exception:
        await status_msg.edit_text("Ошибка поиска.")
    finally:
        await state.set_state(DialogStates.main)


async def process_creative_request(message: Message, state: FSMContext, prompt: str):
    status_msg = await message.answer("🖋️ Генерирую...")
    res = gpt_service.generate_response(prompt, message.text)
    ans = res.get("text", "Ошибка.")
    await status_msg.delete()
    await send_split_message(message, f"<b>Ваш черновик:</b>\n\n{ans}")
    await message.answer("Выберите следующий жанр или выйдите:", reply_markup=create_creative_keyboard())


@router.message(StateFilter(DialogStates.creative_post))
async def gen_post(message: Message, state: FSMContext): await process_creative_request(message, state, POST_PROMPT)


@router.message(StateFilter(DialogStates.creative_release))
async def gen_rel(message: Message, state: FSMContext): await process_creative_request(message, state,
                                                                                       PRESS_RELEASE_PROMPT)


@router.message(StateFilter(DialogStates.creative_announcement))
async def gen_ann(message: Message, state: FSMContext): await process_creative_request(message, state,
                                                                                       ANNOUNCEMENT_PROMPT)


@router.message(StateFilter(DialogStates.creative_custom))
async def gen_cust(message: Message, state: FSMContext): await process_creative_request(message, state,
                                                                                        CUSTOM_CREATIVE_PROMPT)


@router.message(F.photo, StateFilter(DialogStates.recognition_mode))
async def handle_photo_recognition(message: Message, bot: Bot, state: FSMContext):
    fsm_data = await state.get_data()
    recog_type = fsm_data.get("recognition_type", "simple")
    status_msg = await message.reply("⏳ Обрабатываю изображение...")
    try:
        if recog_type == "simple":
            photo_bytes = io.BytesIO()
            await bot.download(message.photo[-1], destination=photo_bytes)
            raw_text = ocr_service.recognize_text(photo_bytes.getvalue())
            if raw_text:
                await status_msg.edit_text("🧹 Чищу текст от мусора...")
                res = gpt_service.generate_response(OCR_CLEANUP_PROMPT, raw_text)
                result_text = res.get("text", raw_text)
            else:
                result_text = None
        elif recog_type == "complex":
            img_base64 = await encode_image_to_base64(bot, message.photo[-1].file_id)
            result_text = await gpt_service.generate_vlm_response(VLM_COMPLEX_PROMPT, img_base64)
            if result_text:
                await status_msg.edit_text("📄 Создаю файл Word...")
                docx_buf = create_formatted_docx(result_text)
                await message.reply_document(BufferedInputFile(docx_buf.getvalue(), filename="document.docx"),
                                             caption="✅ Файл готов.")
                await state.update_data(last_recognized_text=result_text[:3500])
                await status_msg.delete()
                await message.answer("Готов к следующему фото или вопросам:",
                                     reply_markup=create_recognition_keyboard(recog_type))
                return
        elif recog_type == "describe":
            img_base64 = await encode_image_to_base64(bot, message.photo[-1].file_id)
            result_text = await gpt_service.generate_vlm_response(VLM_DESCRIBE_PROMPT, img_base64)

        if result_text:
            await state.update_data(last_recognized_text=result_text[:3500])
            try:
                await status_msg.delete()
            except:
                pass
            await send_split_message(message, f"📄 <b>Результат:</b>\n\n{result_text}")
            await message.answer("Жду следующее фото или вопросы:",
                                 reply_markup=create_recognition_keyboard(recog_type))
        else:
            await status_msg.edit_text("😕 Не удалось распознать.")
    except Exception as e:
        await message.answer(f"⚠️ Ошибка: {e}")


@router.message(F.voice | F.audio, StateFilter(DialogStates.recognition_mode))
async def handle_audio_recognition(message: Message, bot: Bot, state: FSMContext):
    fsm_data = await state.get_data()
    if fsm_data.get("recognition_type") != "audio":
        await message.reply("Выберите режим 'Распознать аудио' в меню параметров.")
        return
    audio_obj = message.voice or message.audio
    if audio_obj and audio_obj.file_size > MAX_AUDIO_SIZE:
        await message.reply("⚠️ Файл слишком большой (лимит 1 МБ).")
        return
    status_msg = await message.reply("⏳ Слушаю аудио...")
    try:
        audio_bytes = io.BytesIO()
        await bot.download(audio_obj, destination=audio_bytes)
        text = await speech_service.speech_to_text(audio_bytes.getvalue())
        if text:
            await state.update_data(last_recognized_text=text[:3500])
            await status_msg.delete()
            await send_split_message(message, f"🎙️ <b>Распознанная речь:</b>\n\n{text}")
            await message.answer("Готов к следующему аудио:", reply_markup=create_recognition_keyboard("audio"))
        else:
            await status_msg.edit_text("😕 Не удалось распознать аудио.")
    except Exception as e:
        await status_msg.edit_text(f"Ошибка аудио: {e}")


@router.message(F.voice)
async def handle_voice_message(message: Message, bot: Bot, state: FSMContext):
    if not message.voice or message.voice.file_size > MAX_AUDIO_SIZE: return
    fsm_data = await state.get_data()
    settings = fsm_data.get("settings", {})
    voice_mode = settings.get("voice_mode", "text_to_text")
    status_msg = await message.reply("🎤 Слушаю...")
    try:
        voice_bytes_io = io.BytesIO()
        await bot.download(message.voice, destination=voice_bytes_io)
        recognized_text = await speech_service.speech_to_text(voice_bytes_io.getvalue())
        if not recognized_text:
            await status_msg.edit_text("😕 Не понял вас.")
            return
        if voice_mode == "voice_to_text":
            await status_msg.edit_text(f"<i>Вы сказали:</i>\n\n{escape(recognized_text)}", parse_mode="HTML")
        elif voice_mode == "voice_to_voice":
            ai_text, _, _, _ = await get_ai_response(state, message.from_user.id, recognized_text)
            voice_res = await speech_service.text_to_speech(ai_text)
            if voice_res:
                await status_msg.delete()
                await message.reply_voice(BufferedInputFile(voice_res, "ans.ogg"))
        else:
            ai_text, suggestions, pdf_slug, metadata = await get_ai_response(state, message.from_user.id,
                                                                             recognized_text)
            final = ai_text + (
                f"\n\n📚 <i>Источник: {escape(str(metadata.get('title')))}</i>" if metadata and metadata.get(
                    'title') else "")
            await status_msg.edit_text(clean_html_for_telegram(final),
                                       reply_markup=create_smart_keyboard(suggestions, pdf_slug), parse_mode="HTML")
    except Exception:
        await status_msg.edit_text("Ошибка голоса.")


@router.message(F.text, StateFilter(None, DialogStates.main, DialogStates.recognition_mode))
async def handle_text_query(message: Message, bot: Bot, state: FSMContext):
    if message.text in ["✍️ Написать нам", "🌐 Поиск в сети", "⚙️ Параметры", "💡 Есть идея",
                        "✨ Креативный режим"]: return
    if await try_send_file(message, bot): return
    fsm_data = await state.get_data()
    settings = fsm_data.get("settings", {})
    voice_mode = settings.get("voice_mode", "text_to_text")
    if voice_mode == "text_playback":
        voice_bytes = await speech_service.text_to_speech(message.text)
        if voice_bytes: await message.reply_voice(BufferedInputFile(voice_bytes, "play.ogg"))
        return
    if voice_mode == "text_to_voice":
        ai_text, _, _, _ = await get_ai_response(state, message.from_user.id, message.text)
        voice_bytes = await speech_service.text_to_speech(ai_text)
        if voice_bytes: await message.reply_voice(BufferedInputFile(voice_bytes, "ans.ogg"))
        return

    # Отправляем "Готовлю ответ" только если это не small_talk, чтобы не моргало лишний раз
    if not is_small_talk(message.text):
        status_msg = await message.answer("💭 Думаю...")
    else:
        status_msg = None

    ai_text, suggestions, pdf_slug, metadata = await get_ai_response(state, message.from_user.id, message.text)
    final_text = ai_text + (
        f"\n\n📚 <i>Источник: {escape(str(metadata.get('title')))}</i>" if metadata and metadata.get('title') else "")

    if status_msg:
        await status_msg.delete()

    await send_split_message(message, final_text, reply_markup=create_smart_keyboard(suggestions, pdf_slug))


@router.callback_query(F.data == "enter_recognition_menu")
async def recognition_menu_handler(callback: CallbackQuery, state: FSMContext):
    await state.set_state(DialogStates.recognition_mode)
    fsm_data = await state.get_data()
    current_type = fsm_data.get("recognition_type", "simple")
    await callback.message.edit_text("<b>Режим распознавания.</b>\nВыберите инструмент:",
                                     reply_markup=create_recognition_keyboard(current_type), parse_mode="HTML")
    await callback.answer()


@router.callback_query(StateFilter(DialogStates.recognition_mode), F.data.startswith("set_recog:"))
async def set_recognition_type(callback: CallbackQuery, state: FSMContext):
    recog_type = callback.data.split(":")[1]
    await state.update_data(recognition_type=recog_type)
    try:
        await callback.message.edit_reply_markup(reply_markup=create_recognition_keyboard(recog_type))
    except:
        pass
    await callback.answer(f"Режим изменен")


@router.callback_query(StateFilter(DialogStates.recognition_mode), F.data == "recog_help")
async def recog_help_handler(callback: CallbackQuery):
    help_text = ("• <b>Простой текст</b> — для обычных страниц книг.\n"
                 "• <b>Сложный документ</b> — для таблиц и документов со сложной версткой.\n"
                 "• <b>Описать изображение</b> — если нужно понять, что на фото.\n"
                 "• <b>Распознать аудио</b> — переведет голосовое или аудиофайл в текст.")
    await callback.message.answer(help_text, parse_mode="HTML")
    await callback.answer()


@router.callback_query(StateFilter(DialogStates.recognition_mode), F.data == "recog_exit")
async def recog_exit_handler(callback: CallbackQuery, state: FSMContext):
    await state.set_state(DialogStates.main)
    await callback.message.edit_text("Вы вышли из режима распознавания.")
    await callback.answer()


@router.callback_query(F.data == "enter_creative_from_settings")
async def enter_creative(callback: CallbackQuery, state: FSMContext):
    await state.set_state(DialogStates.creative_mode)
    await callback.message.edit_text("Креативный режим:", reply_markup=create_creative_keyboard())
    await callback.answer()


@router.callback_query(StateFilter(DialogStates.creative_mode), F.data.startswith("creative:"))
async def select_genre(callback: CallbackQuery, state: FSMContext):
    genre = callback.data.split(":")[1]
    if genre == "exit":
        await state.set_state(DialogStates.main)
        await callback.message.edit_text("Обычный режим.")
    else:
        mapping = {"post": (DialogStates.creative_post, "Тема поста:"),
                   "release": (DialogStates.creative_release, "Событие:"),
                   "announcement": (DialogStates.creative_announcement, "Мероприятие:"),
                   "custom": (DialogStates.creative_custom, "Описание:")}
        new_state, txt = mapping[genre]
        await state.set_state(new_state)
        await callback.message.edit_text(txt)
    await callback.answer()


@router.callback_query(F.data.startswith("ask_suggestion:"))
async def handle_suggestion(callback: CallbackQuery, bot: Bot, state: FSMContext):
    """
    Обработчик нажатия на кнопку-подсказку (Inline-кнопка под ответом).
    """
    try:
        # 1. Извлекаем индекс подсказки из callback_data
        idx = int(callback.data.split(":")[1])

        # 2. Получаем список подсказок из состояния FSM
        data = await state.get_data()
        s_list = data.get("last_suggestions") or STARTUP_SUGGESTIONS

        # Проверяем валидность индекса
        if not (0 <= idx < len(s_list)):
            await callback.answer("Ошибка: подсказка не найдена.")
            return

        txt = s_list[idx]
        await callback.answer()  # Убираем "часики" на кнопке

        # 3. Создаем индикатор ожидания
        status_msg = await callback.message.answer(f"💭 Готовлю ответ на вопрос: «{escape(txt)}»...")
        await bot.send_chat_action(callback.message.chat.id, "typing")

        # 4. Получаем ответ от ИИ (распаковываем 4 значения)
        ai_text, suggestions, pdf_slug, metadata = await get_ai_response(state, callback.from_user.id, txt)

        # 5. Формируем текст ответа с указанием источника
        source_text = ""
        if metadata and metadata.get('title'):
            source_text = f"\n\n📚 <i>Источник: {escape(str(metadata.get('title')))}</i>"

        final_text = ai_text + source_text

        # 6. Очищаем текст от запрещенных HTML-тегов (типа <br>) перед отправкой
        safe_text = clean_html_for_telegram(final_text)

        # 7. Обновляем сообщение (используем клавиатуру с НОВЫМИ подсказками)
        try:
            await status_msg.edit_text(
                safe_text,
                reply_markup=create_smart_keyboard(suggestions, pdf_slug),
                parse_mode="HTML"
            )
        except TelegramBadRequest as e:
            if "message is not modified" in str(e):
                # Если текст идентичен, просто ничего не делаем
                pass
            else:
                # Если HTML всё еще сломан (критическая ошибка парсинга), отправляем без разметки
                logger.error(f"Final HTML error in handle_suggestion: {e}")
                await status_msg.edit_text(escape(safe_text), reply_markup=create_smart_keyboard(suggestions, pdf_slug))

    except (ValueError, IndexError) as e:
        logger.error(f"Ошибка индексации подсказки: {e}")
        await callback.answer("Произошла ошибка при выборе подсказки.", show_alert=True)
    except Exception as e:
        logger.error(f"Критическая ошибка в handle_suggestion: {e}")
        await callback.message.answer("⚠️ Не удалось обработать запрос. Попробуйте написать вопрос вручную.")

@router.callback_query(F.data == "regenerate")
async def handle_regen(callback: CallbackQuery, bot: Bot, state: FSMContext):
    data = await state.get_data()
    last = data.get("last_query")
    if last:
        await callback.answer("Генерирую...")
        status_msg = await callback.message.answer("💭 Переосмысливаю...")
        ai_text, suggestions, pdf_slug, metadata = await get_ai_response(state, callback.from_user.id, last)
        await status_msg.edit_text(ai_text, reply_markup=create_smart_keyboard(suggestions, pdf_slug),
                                   parse_mode="HTML")
    else:
        await callback.answer("Нет запроса.", show_alert=True)


@router.callback_query(F.data.startswith("get_pdf:"))
async def handle_pdf(callback: CallbackQuery, bot: Bot):
    slug = callback.data.split(":")[1]
    fname = rag_service.get_filename_by_slug(slug)
    if fname and (PDF_DIR / fname).exists():
        await bot.send_chat_action(callback.message.chat.id, "upload_document")
        await callback.message.answer_document(FSInputFile(PDF_DIR / fname))
    await callback.answer()


@router.callback_query(F.data.startswith("set_voice_mode:"))
async def set_v_mode(callback: CallbackQuery, state: FSMContext):
    mode = callback.data.split(":")[1]
    data = await state.get_data()
    sets = data.get("settings", {})
    sets["voice_mode"] = mode
    await state.update_data(settings=sets)
    await callback.message.edit_reply_markup(reply_markup=create_settings_keyboard(sets))
    await callback.answer("Режим изменен")


@router.callback_query(F.data == "close_settings")
async def close_sets(callback: CallbackQuery, state: FSMContext):
    await state.set_state(DialogStates.main)
    await callback.message.delete()
    await callback.answer("Сохранено")