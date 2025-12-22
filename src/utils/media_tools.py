import io
import base64
import re
import logging
import qrcode
import cv2
import numpy as np
from typing import Optional
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from aiogram import Bot

# Импорты для распознавания
from pyzbar.pyzbar import decode
from PIL import Image

logger = logging.getLogger(__name__)


async def decode_qr_code(bot: Bot, photo_id: str) -> str:
    """
    Улучшенное распознавание QR-кода с предварительной обработкой изображения.
    Эффективно борется с муаром (сеткой пикселей) при фото с экрана.
    """
    try:
        # 1. Загрузка изображения
        file_info = await bot.get_file(photo_id)
        photo_bytes = io.BytesIO()
        await bot.download_file(file_info.file_path, photo_bytes)

        # Конвертируем байты в формат OpenCV (numpy array)
        file_bytes = np.frombuffer(photo_bytes.getvalue(), dtype=np.uint8)
        image = cv2.imdecode(file_bytes, cv2.IMREAD_COLOR)

        if image is None:
            return ""

        # 2. ПРЕДОБРАБОТКА (Pre-processing)
        # Перевод в градации серого
        gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)

        # Применяем легкое размытие по Гауссу.
        # Это "размывает" пиксельную сетку монитора, помогая pyzbar увидеть общую структуру.
        blurred = cv2.GaussianBlur(gray, (3, 3), 0)

        # Адаптивная бинаризация для создания высокого контраста
        thresh = cv2.threshold(blurred, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)[1]

        # 3. ПОПЫТКА РАСПОЗНАВАНИЯ
        # Сначала пробуем распознать обработанное изображение
        decoded_objects = decode(thresh)

        # Если не вышло, пробуем на исходном сером (иногда фильтры только мешают)
        if not decoded_objects:
            decoded_objects = decode(gray)

        # Если всё еще не вышло, пробуем инвертированное изображение
        if not decoded_objects:
            inverted = cv2.bitwise_not(thresh)
            decoded_objects = decode(inverted)

        if decoded_objects:
            qr_data = decoded_objects[0].data.decode('utf-8')
            logger.info(f"QR успешно распознан после обработки: {qr_data}")
            return qr_data

        logger.warning("QR-код не найден даже после фильтрации изображения.")
        return ""

    except Exception as e:
        logger.error(f"Ошибка в улучшенном декодере QR: {e}")
        return ""


def generate_qr_image(text: str) -> io.BytesIO:
    """Генерация QR-кода (без изменений)."""
    qr = qrcode.QRCode(version=1, box_size=10, border=4)
    qr.add_data(text)
    qr.make(fit=True)
    img = qr.make_image(fill_color="black", back_color="white")
    buf = io.BytesIO()
    img.save(buf, format='PNG')
    buf.seek(0)
    return buf


def create_formatted_docx(md_text: str, title: str = "Распознанный документ") -> io.BytesIO:
    """Создание DOCX (без изменений)."""
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    clean_text = re.sub(r'^(Вот|Here is|Результат|Analysis|Извлеченный|Ниже).*?[:\n]', '', md_text,
                        flags=re.IGNORECASE | re.DOTALL).strip()
    doc.add_heading(title, 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    for line in clean_text.split('\n'):
        line = line.strip()
        if line: doc.add_paragraph(line)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


async def encode_image_to_base64(bot: Bot, photo_id: str) -> str:
    """Кодирование в Base64 (без изменений)."""
    file_info = await bot.get_file(photo_id)
    photo_bytes = io.BytesIO()
    await bot.download_file(file_info.file_path, photo_bytes)
    return base64.b64encode(photo_bytes.getvalue()).decode('utf-8')